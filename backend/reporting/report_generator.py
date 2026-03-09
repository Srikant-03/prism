"""
Report Generator — Builds a comprehensive analysis report from profiling,
cleaning, and insights data. Exports to PDF, DOCX, HTML, and Jupyter Notebook.
"""

from __future__ import annotations

import io
import json
import time
from datetime import datetime
from typing import Any, Optional

import pandas as pd


class ReportSection:
    """One section of the full report."""

    def __init__(self, title: str, content: str, subsections: list[dict] = None,
                 tables: list[dict] = None, charts: list[dict] = None):
        self.title = title
        self.content = content
        self.subsections = subsections or []
        self.tables = tables or []
        self.charts = charts or []

    def to_dict(self) -> dict:
        return {
            "title": self.title,
            "content": self.content,
            "subsections": self.subsections,
            "tables": self.tables,
            "charts": self.charts,
        }


class FullReport:
    """Complete analysis report with all sections."""

    def __init__(self):
        self.title = "Data Intelligence Platform — Full Analysis Report"
        self.generated_at = datetime.utcnow().isoformat()
        self.sections: list[ReportSection] = []

    def add_section(self, section: ReportSection):
        self.sections.append(section)

    def to_dict(self) -> dict:
        return {
            "title": self.title,
            "generated_at": self.generated_at,
            "sections": [s.to_dict() for s in self.sections],
        }


class ReportGenerator:
    """
    Generates a comprehensive analysis report combining:
    - Executive summary
    - Dataset overview & profiling
    - Data quality scores
    - Preprocessing decision log
    - Before/after statistics
    - Feature importance
    - Recommendations
    """

    @staticmethod
    def generate(
        profile_data: Optional[dict] = None,
        cleaning_data: Optional[dict] = None,
        insights_data: Optional[dict] = None,
        audit_log: Optional[list[dict]] = None,
        before_after: Optional[dict] = None,
    ) -> FullReport:
        """Build the full report from all available data."""
        report = FullReport()

        # 1. Executive summary
        report.add_section(ReportGenerator._executive_summary(
            profile_data, insights_data, cleaning_data))

        # 2. Dataset overview
        if profile_data:
            report.add_section(ReportGenerator._dataset_overview(profile_data))

        # 2.5 Data completeness heatmap
        if profile_data and "columns" in profile_data:
            completeness = ReportGenerator._data_completeness(profile_data)
            if completeness:
                report.add_section(completeness)

        # 3. Data quality assessment
        if insights_data and "quality_score" in insights_data:
            report.add_section(ReportGenerator._quality_assessment(
                insights_data["quality_score"]))

        # 4. Profiling findings
        if profile_data and "columns" in profile_data:
            report.add_section(ReportGenerator._profiling_findings(profile_data))

        # 4.5 Numeric distribution analysis
        if profile_data and "columns" in profile_data:
            dist = ReportGenerator._distribution_analysis(profile_data)
            if dist:
                report.add_section(dist)

        # 4.6 Correlation analysis
        if profile_data:
            corr = ReportGenerator._correlation_analysis(profile_data)
            if corr:
                report.add_section(corr)

        # 4.7 Target variable analysis
        if profile_data:
            target_section = ReportGenerator._target_variable_analysis(profile_data)
            if target_section:
                report.add_section(target_section)

        # 5. Preprocessing decisions log
        if audit_log:
            report.add_section(ReportGenerator._preprocessing_log(audit_log))

        # 6. Before/after statistics
        if before_after:
            report.add_section(ReportGenerator._before_after_stats(before_after))

        # 7. Feature importance
        if insights_data and "feature_rankings" in insights_data:
            report.add_section(ReportGenerator._feature_importance(
                insights_data["feature_rankings"]))

        # 8. Anomaly warnings
        if insights_data and "anomalies" in insights_data:
            report.add_section(ReportGenerator._anomaly_summary(
                insights_data["anomalies"]))

        # 8.5 Data-Driven Hypotheses
        if insights_data and "hypotheses" in insights_data and insights_data["hypotheses"]:
            report.add_section(ReportGenerator._hypotheses_insights(
                insights_data["hypotheses"]))

        # 9. Recommendations
        report.add_section(ReportGenerator._recommendations(
            profile_data, insights_data, cleaning_data))

        return report

    # ── Section builders ─────────────────────────────────────────────

    @staticmethod
    def _executive_summary(profile: dict = None, insights: dict = None,
                           cleaning: dict = None) -> ReportSection:
        """Build executive summary."""
        parts = []
        if profile:
            rows = profile.get("total_rows", 0)
            cols = profile.get("total_columns", 0)
            domain = profile.get("estimated_domain", "Unknown")
            parts.append(
                f"This report presents a comprehensive statistical profile, data quality assessment, and exploratory analysis "
                f"for the provided dataset. The underlying data consists of {rows:,} observations across {cols} distinct features, "
                f"and has been algorithmically classified as belonging to the **{domain}** domain."
            )

        if insights and "quality_score" in insights:
            qs = insights["quality_score"]
            grade = qs.get("grade", "N/A")
            score = qs.get("overall_score", 0)
            parts.append(
                f"\n\nBased on evaluating completeness, uniqueness, consistency, and validity dimensions, the overall data quality "
                f"is scored at **{score}/100** (Grade: **{grade}**)."
            )

        if insights and "analyst_briefing" in insights:
            brief = insights["analyst_briefing"]
            if isinstance(brief, dict) and "executive_summary" in brief:
                parts.append(brief["executive_summary"])

        if cleaning:
            total = cleaning.get("total_actions", 0)
            applied = cleaning.get("applied", 0)
            parts.append(
                f"{applied} of {total} recommended preprocessing steps were applied."
            )

        # Dataset readiness assessment
        if profile:
            readiness_issues = []
            columns = profile.get("columns", [])
            if isinstance(columns, dict):
                columns = list(columns.values())
            high_null_cols = sum(1 for c in columns if isinstance(c, dict) and c.get("null_percentage", 0) > 30)
            if high_null_cols:
                readiness_issues.append(f"{high_null_cols} feature(s) with >30% missing values")
            if profile.get("duplicate_row_count", 0) > 0:
                readiness_issues.append(f"{profile['duplicate_row_count']:,} duplicate rows")

            if readiness_issues:
                parts.append(
                    f"\n\n**Dataset Readiness:** The dataset requires preprocessing before modeling. "
                    f"Key issues: {'; '.join(readiness_issues)}."
                )
            else:
                parts.append(
                    "\n\n**Dataset Readiness:** The dataset appears clean and well-structured for immediate exploratory analysis and modeling."
                )

        content = " ".join(parts) if parts else "No data available for summary."
        return ReportSection("Executive Summary", content)

    @staticmethod
    def _dataset_overview(profile: dict) -> ReportSection:
        """Dataset shape, memory, types breakdown."""
        rows = profile.get("total_rows", 0)
        cols = profile.get("total_columns", 0)
        mem = profile.get("memory_size_bytes", 0) / (1024 * 1024)
        domain = profile.get("estimated_domain", "Unknown")
        dup = profile.get("duplicate_row_count", 0)

        content = (
            f"The dataset's memory footprint is approximately **{mem:.2f} MB**. "
            f"It contains **{rows:,}** total observations and **{cols}** distinct features. "
            f"An initial structural scan detected **{dup:,}** completely duplicated rows."
        )

        # Column type breakdown table
        columns = profile.get("columns", [])
        if isinstance(columns, dict):
            columns = list(columns.values())
        type_counts = {}
        for c in columns:
            if isinstance(c, dict):
                t = c.get("semantic_type", c.get("ui_type", "unknown"))
                # Strip 'SemanticType.' prefix for clean display
                t = str(t).replace("SemanticType.", "").lower()
                type_counts[t] = type_counts.get(t, 0) + 1

        # Compute summary stats for intro
        num_numeric = sum(1 for c in columns if isinstance(c, dict) and c.get("inferred_dtype", c.get("dtype", "")) in ("int64", "float64", "int32", "float32"))
        num_cat = sum(1 for c in columns if isinstance(c, dict) and c.get("inferred_dtype", c.get("dtype", "")) in ("object", "category", "string", "bool"))
        avg_null = 0
        if columns:
            avg_null = sum(c.get("null_percentage", 0) for c in columns if isinstance(c, dict)) / len(columns)

        content += (
            f" The feature space is composed of **{num_numeric}** numeric and **{num_cat}** categorical variables. "
            f"Average feature-level missing rate is **{avg_null:.2f}%**."
        )

        tables = [{
            "title": "Feature Type Distribution",
            "headers": ["Semantic Type", "Count"],
            "rows": [[t, str(c)] for t, c in sorted(type_counts.items())],
        }]

        charts = [{
            "type": "pie",
            "title": "Feature Types Breakdown",
            "data": [{"label": str(t), "value": c} for t, c in sorted(type_counts.items())]
        }]

        return ReportSection("Dataset Overview", content, tables=tables, charts=charts)

    @staticmethod
    def _data_completeness(profile: dict) -> ReportSection | None:
        """Data completeness analysis with bar chart."""
        columns = profile.get("columns", [])
        if isinstance(columns, dict):
            columns = list(columns.values())

        completeness_data = []
        for c in columns:
            if not isinstance(c, dict):
                continue
            null_pct = c.get("null_percentage", c.get("null_pct", 0))
            completeness = 100.0 - null_pct
            completeness_data.append({
                "label": c.get("name", "unknown"),
                "value": round(completeness, 1)
            })

        if not completeness_data:
            return None

        # Sort by completeness (lowest first to highlight problems)
        completeness_data.sort(key=lambda x: x["value"])

        fully_complete = sum(1 for d in completeness_data if d["value"] >= 100)
        sparse = sum(1 for d in completeness_data if d["value"] < 70)

        content = (
            f"Out of {len(completeness_data)} features, **{fully_complete}** are fully populated (100% completeness). "
        )
        if sparse > 0:
            content += f"**{sparse}** feature(s) have completeness below 70%, which may require imputation or removal before modeling. "
        else:
            content += "All features exhibit healthy completeness levels above 70%. "

        content += "The chart below visualizes how complete each feature is across the dataset."

        charts = [{
            "type": "bar",
            "title": "Feature Completeness (%)",
            "data": completeness_data[:20]
        }]

        return ReportSection("Data Completeness Analysis", content, charts=charts)

    @staticmethod
    def _quality_assessment(quality: dict) -> ReportSection:
        """Data quality dimension scores."""
        content = (
            f"Overall quality grade: **{quality.get('grade', 'N/A')}** "
            f"(Score: {quality.get('overall_score', 0)}/100)"
        )

        dimensions = [
            ("Completeness", quality.get("completeness", 0)),
            ("Uniqueness", quality.get("uniqueness", 0)),
            ("Validity", quality.get("validity", 0)),
            ("Consistency", quality.get("consistency", 0)),
            ("Timeliness", quality.get("timeliness", 0)),
        ]

        tables = [{
            "title": "Quality Dimensions",
            "headers": ["Dimension", "Score (/100)"],
            "rows": [[d, str(s)] for d, s in dimensions if s],
        }]

        return ReportSection("Data Quality Assessment", content, tables=tables)

    @staticmethod
    def _profiling_findings(profile: dict) -> ReportSection:
        """Per-column profiling stats."""
        columns = profile.get("columns", [])
        if isinstance(columns, dict):
            columns = list(columns.values())

        rows_data = []
        null_data = []
        for c in columns[:50]:
            if not isinstance(c, dict):
                continue
            null_pct = c.get('null_percentage', c.get('null_pct', 0))
            rows_data.append([
                c.get("name", ""),
                c.get("inferred_dtype", c.get("dtype", "")),
                f"{null_pct:.1f}%",
                str(c.get("distinct_count", c.get("unique_count", 0))),
                c.get("semantic_type", ""),
            ])
            if null_pct > 0:
                null_data.append({"label": c.get("name", ""), "value": float(null_pct)})

        tables = [{
            "title": "Feature Catalog",
            "headers": ["Feature", "Data Type", "Missing %", "Distinct Values", "Semantic Type"],
            "rows": rows_data,
        }]

        # Numeric summary sub-table
        numeric_rows = []
        for c in columns[:50]:
            if not isinstance(c, dict):
                continue
            dtype = c.get("inferred_dtype", c.get("dtype", ""))
            if dtype not in ("int64", "float64", "int32", "float32"):
                continue
            num = c.get("numeric") or {}
            if not isinstance(num, dict):
                continue
            numeric_rows.append([
                c.get("name", ""),
                f"{num.get('mean', 0):.2f}",
                f"{num.get('std', 0):.2f}",
                f"{num.get('min', 0):.2f}",
                f"{num.get('percentile_25', num.get('q1', 0)):.2f}",
                f"{num.get('median', num.get('percentile_50', 0)):.2f}",
                f"{num.get('percentile_75', num.get('q3', 0)):.2f}",
                f"{num.get('max', 0):.2f}",
            ])

        if numeric_rows:
            tables.append({
                "title": "Descriptive Statistics (Numeric Features)",
                "headers": ["Feature", "Mean", "Std Dev", "Min", "Q1 (25%)", "Median", "Q3 (75%)", "Max"],
                "rows": numeric_rows,
            })

        charts = []
        if null_data:
            null_data.sort(key=lambda x: x["value"], reverse=True)
            charts.append({
                "type": "bar",
                "title": "Missing Values (%) by Feature",
                "data": null_data[:15]
            })

        content = f"The following table summarizes the structural data types, sparsity (missing value percentages), and cardinality (distinct values) for all {len(columns)} features."
        if numeric_rows:
            content += f" A separate descriptive statistics table is included for {len(numeric_rows)} numeric features, showing central tendency (mean, median) and spread (std, quartiles)."

        return ReportSection("Statistical Profiling & Feature Distributions", content, tables=tables, charts=charts)

    @staticmethod
    def _distribution_analysis(profile: dict) -> ReportSection | None:
        """Analyze distributions of numeric columns — skewness, kurtosis."""
        columns = profile.get("columns", [])
        if isinstance(columns, dict):
            columns = list(columns.values())

        dist_rows = []
        skew_chart = []
        for c in columns:
            if not isinstance(c, dict):
                continue
            dtype = c.get("inferred_dtype", c.get("dtype", ""))
            if dtype not in ("int64", "float64", "int32", "float32"):
                continue
            num = c.get("numeric") or {}
            if not isinstance(num, dict):
                continue
            skew = num.get("skewness")
            kurt = num.get("kurtosis")
            if skew is None and kurt is None:
                continue

            skew_val = skew if skew is not None else 0
            kurt_val = kurt if kurt is not None else 0

            # Interpret skewness
            if abs(skew_val) < 0.5:
                skew_interp = "Approximately symmetric"
            elif skew_val > 0:
                skew_interp = f"Right-skewed ({skew_val:.2f})"
            else:
                skew_interp = f"Left-skewed ({skew_val:.2f})"

            # Interpret kurtosis
            if abs(kurt_val) < 1:
                kurt_interp = "Mesokurtic (normal-like)"
            elif kurt_val > 0:
                kurt_interp = f"Leptokurtic/heavy-tailed ({kurt_val:.2f})"
            else:
                kurt_interp = f"Platykurtic/light-tailed ({kurt_val:.2f})"

            dist_rows.append([
                c.get("name", ""),
                f"{skew_val:.3f}",
                skew_interp,
                f"{kurt_val:.3f}",
                kurt_interp,
            ])
            skew_chart.append({"label": c.get("name", ""), "value": round(abs(skew_val), 2)})

        if not dist_rows:
            return None

        skew_chart.sort(key=lambda x: x["value"], reverse=True)

        tables = [{
            "title": "Distribution Shape Analysis",
            "headers": ["Feature", "Skewness", "Interpretation", "Kurtosis", "Interpretation"],
            "rows": dist_rows,
        }]

        charts = [{
            "type": "bar",
            "title": "Absolute Skewness by Feature (higher = more asymmetric)",
            "data": skew_chart[:15]
        }]

        heavily_skewed = sum(1 for r in dist_rows if abs(float(r[1])) > 2)
        content = (
            f"Distribution shape analysis for {len(dist_rows)} numeric features. "
        )
        if heavily_skewed:
            content += (
                f"**{heavily_skewed}** feature(s) exhibit severe skewness (|skew| > 2), indicating "
                f"a long tail that may distort linear model assumptions. "
                f"Consider log or Box-Cox transformations for these features before fitting parametric models."
            )
        else:
            content += "All numeric features exhibit moderate or near-symmetric distributions."

        return ReportSection("Distribution Shape Analysis", content, tables=tables, charts=charts)

    @staticmethod
    def _correlation_analysis(profile: dict) -> ReportSection | None:
        """Analyze correlations from the profile."""
        cross = profile.get("cross_analysis") or {}
        correlations = cross.get("correlations") or profile.get("correlations") or {}

        if isinstance(correlations, dict) and correlations:
            corr_rows = []
            strong_pairs = []
            for pair, val in correlations.items():
                if not isinstance(val, (int, float)):
                    continue
                strength = "Strong" if abs(val) > 0.7 else ("Moderate" if abs(val) > 0.4 else "Weak")
                direction = "Positive" if val > 0 else "Negative"
                corr_rows.append([str(pair), f"{val:.3f}", direction, strength])
                if abs(val) > 0.7:
                    strong_pairs.append({"label": str(pair)[:30], "value": round(abs(val), 3)})

            if corr_rows:
                corr_rows.sort(key=lambda r: abs(float(r[1])), reverse=True)

                tables = [{
                    "title": "Feature Correlation Matrix (Top Pairs)",
                    "headers": ["Feature Pair", "Correlation (r)", "Direction", "Strength"],
                    "rows": corr_rows[:20],
                }]

                charts = []
                if strong_pairs:
                    strong_pairs.sort(key=lambda x: x["value"], reverse=True)
                    charts.append({
                        "type": "bar",
                        "title": "Strongly Correlated Feature Pairs (|r| > 0.7)",
                        "data": strong_pairs[:10]
                    })

                strong_count = sum(1 for r in corr_rows if r[3] == "Strong")
                content = (
                    f"Pairwise Pearson correlation analysis identified **{len(corr_rows)}** feature pairs. "
                    f"**{strong_count}** pair(s) exhibit strong correlation (|r| > 0.7), suggesting potential "
                    f"multicollinearity. These should be examined before training regression or gradient-based models "
                    f"to avoid inflated coefficient variance."
                )
                return ReportSection("Correlation & Multicollinearity Analysis", content, tables=tables, charts=charts)

        return None

    @staticmethod
    def _target_variable_analysis(profile: dict) -> ReportSection | None:
        """Analyze the detected target variable and its relationship with other features."""
        cross = profile.get("cross_analysis") or {}
        if isinstance(cross, dict):
            target_info = cross.get("target_analysis") or {}
        else:
            try:
                target_info = cross.target_analysis.__dict__ if hasattr(cross, "target_analysis") and cross.target_analysis else {}
            except Exception:
                target_info = {}

        if not isinstance(target_info, dict):
            try:
                target_info = target_info.__dict__
            except Exception:
                return None

        target_col = target_info.get("target_column")
        if not target_col:
            return None

        is_detected = target_info.get("is_target_detected", False)
        if not is_detected:
            return None

        confidence = target_info.get("confidence", 0)
        problem_type = target_info.get("problem_type", "unknown")
        justification = target_info.get("justification", "")
        class_dist = target_info.get("class_distribution") or {}
        imbalance = target_info.get("imbalance_ratio")
        top_predictors = target_info.get("top_predictors") or []

        # Format problem type for display
        problem_display = problem_type.replace("_", " ").title()

        content = (
            f"The system has automatically identified **'{target_col}'** as the most likely target "
            f"(dependent) variable for predictive modeling. "
            f"The detected problem type is **{problem_display}** with a detection confidence "
            f"of **{confidence:.0%}**.\n\n"
        )

        if justification:
            content += f"**Rationale:** {justification}\n\n"

        if problem_type == "regression":
            content += (
                f"Since '{target_col}' is a continuous numeric variable, this is a **regression** problem. "
                f"Models like Linear Regression, Random Forest Regressor, or Gradient Boosting can be applied. "
                f"Feature selection based on the predictors below will help improve model performance."
            )
        elif "classification" in problem_type:
            n_classes = len(class_dist) if class_dist else "unknown"
            content += (
                f"This is a **classification** problem with {n_classes} distinct classes. "
            )
            if imbalance and imbalance > 3:
                content += (
                    f"The class imbalance ratio is **{imbalance:.1f}:1**, which may cause "
                    f"the model to favor the majority class. Consider SMOTE, class weights, "
                    f"or undersampling strategies."
                )

        tables = []
        charts = []

        # Top predictors table and chart
        if top_predictors:
            pred_rows = []
            pred_chart = []
            for p in top_predictors[:15]:
                if isinstance(p, dict):
                    feat = p.get("feature", "")
                    score = p.get("importance_score", 0)
                else:
                    feat = getattr(p, "feature", "")
                    score = getattr(p, "importance_score", 0)

                # Interpret impact strength
                if score > 0.7:
                    strength = "Very Strong"
                elif score > 0.4:
                    strength = "Moderate"
                elif score > 0.1:
                    strength = "Weak"
                else:
                    strength = "Negligible"

                pred_rows.append([feat, f"{score:.4f}", strength])
                pred_chart.append({"label": feat, "value": round(float(score), 4)})

            tables.append({
                "title": f"Features Most Influencing '{target_col}'",
                "headers": ["Feature", "Association Score", "Impact Strength"],
                "rows": pred_rows,
            })

            pred_chart.sort(key=lambda x: x["value"], reverse=True)
            charts.append({
                "type": "bar",
                "title": f"Top Predictors for '{target_col}' (by association strength)",
                "data": pred_chart[:10]
            })

            # Add analytical narrative about top predictors
            if pred_rows:
                top_feat = pred_rows[0][0]
                top_score = pred_rows[0][1]
                content += (
                    f"\n\n**Key Finding:** The strongest predictor of '{target_col}' is "
                    f"**'{top_feat}'** with an association score of {top_score}. "
                )
                strong_features = [r[0] for r in pred_rows if r[2] in ("Very Strong", "Moderate")]
                if len(strong_features) > 1:
                    content += (
                        f"Features with moderate-to-strong predictive power include: "
                        f"**{', '.join(strong_features[:5])}**. "
                        f"These should be prioritized in any feature engineering or model training pipeline."
                    )

        # Class distribution chart for classification
        if class_dist and isinstance(class_dist, dict):
            dist_chart = [{"label": str(k), "value": round(float(v) * 100, 1)} for k, v in class_dist.items()]
            charts.append({
                "type": "pie",
                "title": f"Class Distribution of '{target_col}'",
                "data": dist_chart[:10]
            })

        return ReportSection("Target Variable & Predictive Analysis", content, tables=tables, charts=charts)

    @staticmethod
    def _preprocessing_log(audit_log: list[dict]) -> ReportSection:
        """Detailed log of every preprocessing decision."""
        rows_data = []
        for entry in audit_log:
            rows_data.append([
                entry.get("step_name", ""),
                entry.get("action_type", ""),
                entry.get("status", ""),
                entry.get("trigger_reason", ""),
                ", ".join(entry.get("columns_affected", [])[:3]),
                f"{entry.get('rows_before', 0)} → {entry.get('rows_after', 0)}",
            ])

        tables = [{
            "title": "Preprocessing Steps",
            "headers": ["Step", "Action", "Status", "Reason", "Columns", "Rows (Before→After)"],
            "rows": rows_data,
        }]

        applied = sum(1 for e in audit_log if e.get("status") == "applied")
        skipped = sum(1 for e in audit_log if e.get("status") == "skipped")

        content = (
            f"Total preprocessing steps: **{len(audit_log)}** "
            f"({applied} applied, {skipped} skipped)"
        )
        return ReportSection("Preprocessing Decisions Log", content, tables=tables)

    @staticmethod
    def _before_after_stats(comparison: dict) -> ReportSection:
        """Before/after statistics for transformed columns."""
        content = (
            f"Shape: {comparison.get('original_shape', '?')} → "
            f"{comparison.get('current_shape', '?')}"
        )

        col_changes = comparison.get("column_changes", {})
        rows_data = []
        for col, changes in col_changes.items():
            if isinstance(changes, dict):
                rows_data.append([
                    col,
                    str(changes.get("null_before", "")),
                    str(changes.get("null_after", "")),
                    str(changes.get("dtype_before", "")),
                    str(changes.get("dtype_after", "")),
                ])

        tables = []
        if rows_data:
            tables.append({
                "title": "Column Transformations",
                "headers": ["Column", "Nulls Before", "Nulls After", "Type Before", "Type After"],
                "rows": rows_data[:30],
            })

        return ReportSection("Before/After Statistics", content, tables=tables)

    @staticmethod
    def _feature_importance(rankings: list) -> ReportSection:
        """Feature importance table."""
        rows_data = []
        chart_data = []
        for r in rankings[:20]:
            if isinstance(r, dict):
                score = r.get("importance_score", 0)
                feat = r.get("feature", "")
                rows_data.append([
                    feat,
                    f"{score:.1f}",
                    r.get("method", ""),
                    r.get("rationale", "")[:80],
                ])
                chart_data.append({"label": feat, "value": float(score)})

        tables = [{
            "title": "Feature Rankings",
            "headers": ["Feature", "Score", "Method", "Rationale"],
            "rows": rows_data,
        }]
        
        charts = [{
            "type": "bar",
            "title": "Top Feature Importance Scores",
            "data": chart_data[:10]
        }] if chart_data else []

        content = f"Top {len(rows_data)} features by predictive importance."
        return ReportSection("Feature Importance", content, tables=tables, charts=charts)

    @staticmethod
    def _anomaly_summary(anomalies: list) -> ReportSection:
        """Anomaly warnings summary."""
        if not anomalies:
            return ReportSection("Anomaly Summary", "No anomalies detected.")

        rows_data = []
        for a in anomalies:
            if isinstance(a, dict):
                rows_data.append([
                    a.get("feature", ""),
                    a.get("severity", ""),
                    a.get("anomaly_type", ""),
                    a.get("description", "")[:100],
                ])

        tables = [{
            "title": "Anomaly Warnings",
            "headers": ["Feature", "Severity", "Type", "Description"],
            "rows": rows_data,
        }]

        critical = sum(1 for a in anomalies if isinstance(a, dict)
                       and a.get("severity", "").lower() in ("critical", "high"))
        content = (
            f"{len(anomalies)} anomalies detected ({critical} critical/high severity)."
        )
        return ReportSection("Anomaly Summary", content, tables=tables)

    @staticmethod
    def _hypotheses_insights(hypotheses: list) -> ReportSection:
        """Strategic analytical hypotheses with both technical and plain-English explanations."""
        if not hypotheses:
            return ReportSection("Deep Dataset Insights",
                                 "No strategic hypotheses could be generated from the available profiling data.")

        rows_data = []
        for h in hypotheses[:12]:
            if isinstance(h, dict):
                rows_data.append([
                    h.get("observation", ""),
                    h.get("evidence", ""),
                    h.get("impact", "").capitalize(),
                    f"{h.get('confidence', 0):.2f}",
                    h.get("question", "")
                ])

        tables = [{
            "title": "Strategic Hypotheses & Modeling Insights",
            "headers": ["Observation", "Evidence", "Impact", "Confidence", "Recommended Action"],
            "rows": rows_data,
        }]

        # Build layman summary table
        layman_rows = []
        for h in hypotheses[:12]:
            if isinstance(h, dict) and h.get("layman"):
                layman_rows.append([
                    h.get("observation", "")[:80] + ("..." if len(h.get("observation", "")) > 80 else ""),
                    h.get("layman", ""),
                ])

        if layman_rows:
            tables.append({
                "title": "Plain-English Explanations",
                "headers": ["Finding", "What This Means"],
                "rows": layman_rows,
            })

        content = (
            "The following strategic insights were automatically derived by analyzing "
            "statistical relationships, distributional properties, and feature interactions "
            "within the dataset. Unlike basic data quality checks, these hypotheses focus "
            "on **modeling strategy** — the kinds of patterns that determine which algorithms "
            "will succeed, which features to engineer, and where hidden risks lie."
        )
        return ReportSection("Deep Dataset Insights", content, tables=tables)

    @staticmethod
    def _recommendations(profile: dict = None, insights: dict = None,
                         cleaning: dict = None) -> ReportSection:
        """Domain-aware, algorithm-aware recommended next steps."""
        recs = []

        # Pull analyst briefing recommendations if available
        if insights and "analyst_briefing" in insights:
            brief = insights["analyst_briefing"]
            if isinstance(brief, dict):
                for action in brief.get("recommended_actions", []):
                    recs.append(action)

        if profile:
            columns = profile.get("columns", [])
            if isinstance(columns, dict):
                columns = list(columns.values())
            total_rows = profile.get("total_rows", 0)
            total_cols = profile.get("total_columns", 0)

            # ── Missing value handling ──
            high_null = [c.get("name", "") for c in columns if isinstance(c, dict) and c.get("null_percentage", 0) > 50]
            moderate_null = [c.get("name", "") for c in columns if isinstance(c, dict) and 5 < c.get("null_percentage", 0) <= 50]
            if high_null:
                recs.append(
                    f"Features with >50% missingness ({', '.join(high_null[:5])}) should be evaluated for removal. "
                    f"If they contain critical domain signals, consider advanced imputation (KNN Imputer, IterativeImputer/MICE)."
                )
            if moderate_null:
                recs.append(
                    f"Features with 5-50% missingness ({', '.join(moderate_null[:5])}) can be imputed using "
                    f"median (numeric) or mode (categorical) strategies, or more sophisticated methods like KNNImputer."
                )

            # ── Skewness handling ──
            skewed_cols = []
            for c in columns:
                if not isinstance(c, dict):
                    continue
                num = c.get("numeric") or {}
                if isinstance(num, dict) and num.get("skewness") is not None and abs(num["skewness"]) > 2:
                    skewed_cols.append(c.get("name", ""))
            if skewed_cols:
                recs.append(
                    f"Highly skewed features ({', '.join(skewed_cols[:5])}) should be transformed using "
                    f"log, sqrt, or Box-Cox transforms before training linear models, SVMs, or KNN."
                )

            # ── High cardinality categoricals ──
            high_card = [c.get("name", "") for c in columns 
                         if isinstance(c, dict) and c.get("inferred_dtype", c.get("dtype", "")) in ("object", "string", "category")
                         and c.get("distinct_count", c.get("unique_count", 0)) > 50]
            if high_card:
                recs.append(
                    f"High-cardinality categorical features ({', '.join(high_card[:5])}) will cause dimensionality explosion with one-hot encoding. "
                    f"Consider Target Encoding, Frequency Encoding, or native handling via CatBoost/LightGBM."
                )

            # ── Dimensionality concerns ──
            if total_cols > 50:
                recs.append(
                    f"With {total_cols} features, consider dimensionality reduction: PCA for numeric features, "
                    f"feature selection via mutual information or L1 regularization, or tree-based feature importance filtering."
                )

            # ── Small dataset warning ──
            if total_rows < 1000:
                recs.append(
                    f"With only {total_rows:,} observations, overfitting risk is high. Use stratified k-fold cross-validation (k=5 or 10), "
                    f"avoid complex models without regularization, and consider data augmentation if applicable."
                )

            # ── Target-specific ML algorithm recommendations ──
            cross = profile.get("cross_analysis", {})
            if isinstance(cross, dict):
                target = cross.get("target_analysis", {})
                if isinstance(target, dict) and target.get("target_column"):
                    target_col = target["target_column"]
                    problem_type = target.get("problem_type", "")
                    imbalance = target.get("imbalance_ratio")

                    if "regression" in problem_type:
                        recs.append(
                            f"For predicting '{target_col}' (regression), recommended algorithms: "
                            f"Linear Regression / Ridge / Lasso (baseline), Random Forest Regressor, "
                            f"XGBoost / LightGBM / CatBoost (gradient boosting), and SVR for non-linear patterns. "
                            f"Evaluate using RMSE, MAE, and R-squared."
                        )
                    elif "binary" in problem_type:
                        recs.append(
                            f"For predicting '{target_col}' (binary classification), recommended algorithms: "
                            f"Logistic Regression (baseline), Random Forest, XGBoost / LightGBM, SVM, "
                            f"and Neural Networks for complex feature interactions. "
                            f"Evaluate using AUC-ROC, F1-Score, Precision, and Recall."
                        )
                    elif "multiclass" in problem_type:
                        recs.append(
                            f"For predicting '{target_col}' (multiclass classification), recommended algorithms: "
                            f"Multinomial Logistic Regression (baseline), Random Forest, XGBoost with multi:softmax, "
                            f"CatBoost (handles categoricals natively), and Neural Networks. "
                            f"Evaluate using macro/weighted F1-Score and confusion matrix."
                        )

                    # Imbalance-specific advice
                    if imbalance and imbalance > 3:
                        recs.append(
                            f"Class imbalance ratio is {imbalance:.1f}:1. Apply SMOTE / ADASYN for oversampling, "
                            f"or use class_weight='balanced' in sklearn models. For gradient boosting, "
                            f"use scale_pos_weight (XGBoost) or is_unbalance (LightGBM)."
                        )

            # ── Correlation-based recommendations ──
            if isinstance(cross, dict):
                corrs = cross.get("correlations") or {}
                if isinstance(corrs, dict):
                    strong_corrs = corrs.get("strongest_pairs") or []
                    very_strong = [p for p in strong_corrs if isinstance(p, dict) and abs(p.get("score", 0)) > 0.9]
                    if very_strong:
                        recs.append(
                            f"{len(very_strong)} feature pair(s) exhibit near-perfect correlation (|r| > 0.9). "
                            f"Use VIF (Variance Inflation Factor) analysis to identify and remove redundant features "
                            f"before training regression or gradient-based models."
                        )

        if not recs:
            recs.append("Dataset appears analysis-ready. Proceed with exploratory data analysis, feature engineering, and model selection.")

        content = "\n".join(f"- {r}" for r in recs)
        return ReportSection("Recommended Next Steps", content)


# ── Export formatters ─────────────────────────────────────────────────

class ReportExporter:
    """Export FullReport to various formats."""

    @staticmethod
    def to_html(report: FullReport) -> str:
        """Generate an HTML report with embedded CSS charts."""
        lines = [
            "<!DOCTYPE html>",
            '<html lang="en"><head>',
            '<meta charset="utf-8"/>',
            f"<title>{report.title}</title>",
            "<style>",
            "body { font-family: 'Segoe UI', sans-serif; max-width: 900px; margin: 40px auto; padding: 20px; background: #0d1117; color: #c9d1d9; }",
            "h1 { color: #58a6ff; border-bottom: 2px solid #21262d; padding-bottom: 10px; }",
            "h2 { color: #79c0ff; margin-top: 30px; }",
            "h3 { color: #8b949e; margin-top: 16px; }",
            "table { border-collapse: collapse; width: 100%; margin: 12px 0; }",
            "th { background: #161b22; color: #58a6ff; padding: 8px 12px; text-align: left; border: 1px solid #30363d; font-size: 12px; }",
            "td { padding: 6px 12px; border: 1px solid #30363d; font-size: 12px; }",
            "tr:nth-child(even) { background: rgba(255,255,255,0.02); }",
            "p { line-height: 1.6; }",
            ".meta { color: #8b949e; font-size: 12px; }",
            "strong { color: #f0f6fc; }",
            # CSS chart styles
            ".chart-container { margin: 16px 0; padding: 16px; background: #161b22; border-radius: 8px; border: 1px solid #30363d; }",
            ".chart-title { font-size: 13px; font-weight: 600; color: #79c0ff; margin-bottom: 12px; }",
            ".bar-chart { display: flex; flex-direction: column; gap: 6px; }",
            ".bar-row { display: flex; align-items: center; gap: 8px; }",
            ".bar-label { min-width: 120px; font-size: 11px; text-align: right; color: #8b949e; }",
            ".bar-track { flex: 1; height: 20px; background: #21262d; border-radius: 4px; overflow: hidden; }",
            ".bar-fill { height: 100%; border-radius: 4px; transition: width 0.3s; display: flex; align-items: center; padding-left: 6px; font-size: 10px; color: white; font-weight: 600; }",
            ".bar-value { min-width: 60px; font-size: 11px; color: #c9d1d9; font-family: monospace; }",
            ".pie-chart { display: flex; align-items: center; gap: 24px; }",
            ".pie-circle { width: 120px; height: 120px; border-radius: 50%; position: relative; }",
            ".pie-legend { display: flex; flex-direction: column; gap: 4px; }",
            ".pie-legend-item { display: flex; align-items: center; gap: 6px; font-size: 11px; }",
            ".pie-legend-dot { width: 10px; height: 10px; border-radius: 2px; }",
            "</style>",
            "</head><body>",
            f"<h1>{report.title}</h1>",
            f'<p class="meta">Generated: {report.generated_at}</p>',
        ]

        chart_colors = ['#56B4E9', '#E69F00', '#009E73', '#CC79A7', '#0072B2', '#D55E00', '#F0E442', '#999']

        for section in report.sections:
            lines.append(f"<h2>{section.title}</h2>")
            # Convert markdown bold to HTML
            content = section.content.replace("**", "<strong>", 1)
            while "**" in content:
                content = content.replace("**", "</strong>", 1)
                if "**" in content:
                    content = content.replace("**", "<strong>", 1)
            lines.append(f"<p>{content}</p>")

            for table in section.tables:
                lines.append(f"<h3>{table.get('title', '')}</h3>")
                lines.append("<table>")
                lines.append("<tr>" + "".join(f"<th>{h}</th>" for h in table.get("headers", [])) + "</tr>")
                for row in table.get("rows", []):
                    lines.append("<tr>" + "".join(f"<td>{v}</td>" for v in row) + "</tr>")
                lines.append("</table>")

            # Render embedded charts
            for chart in section.charts:
                chart_type = chart.get("type", "bar")
                title = chart.get("title", "")
                data = chart.get("data", [])

                lines.append('<div class="chart-container">')
                lines.append(f'<div class="chart-title">{title}</div>')

                if chart_type == "bar" and data:
                    max_val = max((abs(d.get("value", 0)) for d in data), default=1) or 1
                    lines.append('<div class="bar-chart">')
                    for i, d in enumerate(data[:20]):
                        val = d.get("value", 0)
                        pct = min(abs(val) / max_val * 100, 100)
                        color = chart_colors[i % len(chart_colors)]
                        label = str(d.get("label", ""))[:30]
                        lines.append(f'<div class="bar-row">')
                        lines.append(f'  <span class="bar-label">{label}</span>')
                        lines.append(f'  <div class="bar-track"><div class="bar-fill" style="width:{pct:.0f}%;background:{color}">{val}</div></div>')
                        lines.append(f'</div>')
                    lines.append('</div>')

                elif chart_type == "pie" and data:
                    total = sum(d.get("value", 0) for d in data) or 1
                    # Build conic-gradient
                    gradients = []
                    cumulative = 0
                    for i, d in enumerate(data[:8]):
                        val = d.get("value", 0)
                        pct = val / total * 100
                        color = chart_colors[i % len(chart_colors)]
                        gradients.append(f"{color} {cumulative:.1f}% {cumulative + pct:.1f}%")
                        cumulative += pct
                    gradient = ", ".join(gradients)
                    lines.append('<div class="pie-chart">')
                    lines.append(f'  <div class="pie-circle" style="background:conic-gradient({gradient})"></div>')
                    lines.append('  <div class="pie-legend">')
                    for i, d in enumerate(data[:8]):
                        color = chart_colors[i % len(chart_colors)]
                        label = str(d.get("label", ""))[:30]
                        pct = d.get("value", 0) / total * 100
                        lines.append(f'    <div class="pie-legend-item"><div class="pie-legend-dot" style="background:{color}"></div>{label}: {pct:.1f}%</div>')
                    lines.append('  </div>')
                    lines.append('</div>')

                lines.append('</div>')

        lines.append("</body></html>")
        return "\n".join(lines)

    @staticmethod
    def to_pdf(report: FullReport) -> bytes:
        """Generate a PDF report."""
        try:
            from fpdf import FPDF
        except ImportError:
            return b""

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        def sanitize_text(text: str) -> str:
            if not isinstance(text, str):
                text = str(text)
            # Replace common unsupported FPDF unicode chars
            return text.replace("—", "-").replace("×", "x").replace("…", "...").replace("’", "'").replace("‘", "'").replace("•", "-").replace("“", '"').replace("”", '"')

        # Title
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 12, sanitize_text(report.title), ln=True, align="C")
        pdf.set_font("Arial", "", 10)
        pdf.cell(0, 6, sanitize_text(f"Generated: {report.generated_at}"), ln=True, align="C")
        pdf.ln(8)

        for section in report.sections:
            pdf.set_font("Arial", "B", 13)
            pdf.cell(0, 10, sanitize_text(section.title), ln=True)
            pdf.set_font("Arial", "", 10)
            # Clean markdown bold markers
            text = section.content.replace("**", "")
            pdf.multi_cell(0, 5, sanitize_text(text))
            pdf.ln(3)

            for table in section.tables:
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 7, sanitize_text(table.get("title", "")), ln=True)

                headers = table.get("headers", [])
                rows_data = table.get("rows", [])
                if not headers:
                    continue

                col_w = min(180 // len(headers), 50)
                pdf.set_font("Arial", "B", 8)
                for h in headers:
                    pdf.cell(col_w, 6, sanitize_text(str(h)[:20]), border=1)
                pdf.ln()

                pdf.set_font("Arial", "", 8)
                for row in rows_data[:30]:
                    for val in row:
                        pdf.cell(col_w, 5, sanitize_text(str(val)[:20]), border=1)
                    pdf.ln()
                pdf.ln(3)

        return bytes(pdf.output())

    @staticmethod
    def to_docx(report: FullReport) -> bytes:
        """Generate a DOCX report."""
        try:
            from docx import Document
        except ImportError:
            return b""

        doc = Document()
        doc.add_heading(report.title, 0)
        doc.add_paragraph(f"Generated: {report.generated_at}")

        for section in report.sections:
            doc.add_heading(section.title, level=1)
            doc.add_paragraph(section.content.replace("**", ""))

            for table in section.tables:
                doc.add_heading(table.get("title", ""), level=2)
                headers = table.get("headers", [])
                rows_data = table.get("rows", [])
                if not headers or not rows_data:
                    continue

                t = doc.add_table(rows=1, cols=len(headers))
                t.style = "Table Grid"
                for i, h in enumerate(headers):
                    t.rows[0].cells[i].text = str(h)
                for row in rows_data[:30]:
                    cells = t.add_row().cells
                    for i, val in enumerate(row):
                        if i < len(cells):
                            cells[i].text = str(val)

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    @staticmethod
    def to_notebook(report: FullReport) -> str:
        """Generate a Jupyter Notebook (.ipynb) report."""
        cells = []

        # Title cell
        cells.append({
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                f"# {report.title}\n",
                f"*Generated: {report.generated_at}*\n",
            ],
        })

        for section in report.sections:
            # Section header
            cells.append({
                "cell_type": "markdown",
                "metadata": {},
                "source": [f"## {section.title}\n\n{section.content}\n"],
            })

            # Tables as markdown
            for table in section.tables:
                headers = table.get("headers", [])
                rows_data = table.get("rows", [])
                if not headers:
                    continue

                lines = [f"### {table.get('title', '')}\n\n"]
                lines.append("| " + " | ".join(headers) + " |\n")
                lines.append("| " + " | ".join(["---"] * len(headers)) + " |\n")
                for row in rows_data[:30]:
                    lines.append("| " + " | ".join(str(v) for v in row) + " |\n")
                lines.append("\n")

                cells.append({
                    "cell_type": "markdown",
                    "metadata": {},
                    "source": lines,
                })

        notebook = {
            "nbformat": 4,
            "nbformat_minor": 5,
            "metadata": {
                "kernelspec": {
                    "display_name": "Python 3",
                    "language": "python",
                    "name": "python3",
                },
            },
            "cells": cells,
        }

        return json.dumps(notebook, indent=2)

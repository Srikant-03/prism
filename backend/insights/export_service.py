import io
from docx import Document
from fpdf import FPDF
from insights.insight_models import DatasetInsights

class ExportService:
    """
    Handles generation of PDF and DOCX files from the structured DatasetInsights.
    """

    @staticmethod
    def generate_pdf(insights: DatasetInsights) -> bytes:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, "Data Intelligence System - Analyst Briefing", ln=True, align='C')
        pdf.ln(5)

        briefing = insights.analyst_briefing

        # Header 1: Executive Summary
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "1. Executive Summary", ln=True)
        pdf.set_font("Arial", '', 11)
        pdf.multi_cell(0, 6, briefing.executive_summary)
        pdf.ln(3)

        # Header 2: Dataset Characteristics
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "2. Dataset Characteristics", ln=True)
        pdf.set_font("Arial", '', 11)
        pdf.multi_cell(0, 6, briefing.dataset_characteristics)
        pdf.ln(3)

        # Header 3: Quality Assessment
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "3. Quality Assessment", ln=True)
        pdf.set_font("Arial", '', 11)
        pdf.multi_cell(0, 6, briefing.quality_assessment)
        pdf.ln(3)

        # Header 4: Key Findings
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "4. Key Findings", ln=True)
        pdf.set_font("Arial", '', 11)
        for finding in briefing.key_findings:
            pdf.multi_cell(0, 6, f"- {finding}")
        pdf.ln(3)

        # Header 5: Recommended Actions
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, "5. Recommended Actions", ln=True)
        pdf.set_font("Arial", '', 11)
        for action in briefing.recommended_actions:
            pdf.multi_cell(0, 6, f"- {action}")
        
        return bytes(pdf.output())

    @staticmethod
    def generate_docx(insights: DatasetInsights) -> bytes:
        doc = Document()
        doc.add_heading('Data Intelligence System - Analyst Briefing', 0)

        briefing = insights.analyst_briefing

        doc.add_heading('1. Executive Summary', level=1)
        doc.add_paragraph(briefing.executive_summary)

        doc.add_heading('2. Dataset Characteristics', level=1)
        doc.add_paragraph(briefing.dataset_characteristics)

        doc.add_heading('3. Quality Assessment', level=1)
        doc.add_paragraph(briefing.quality_assessment)

        doc.add_heading('4. Key Findings', level=1)
        for finding in briefing.key_findings:
            doc.add_paragraph(finding, style='List Bullet')

        doc.add_heading('5. Recommended Actions', level=1)
        for action in briefing.recommended_actions:
            doc.add_paragraph(action, style='List Bullet')

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()
